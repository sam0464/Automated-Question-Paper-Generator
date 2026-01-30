from docx import Document

def create_question_paper_structure(output_path, questions):
    # Create a new Document
    doc = Document()

    # Title Section
    doc.add_heading('Vishwakarma Institute of Technology', level=1)
    doc.add_paragraph('Title: Question Paper                                                      FF No. 868')
    doc.add_paragraph('Reg. No.                                                    ')
    doc.add_paragraph(
        'Bansilal Ramnath Agarwal Charitable Trust’s\n'
        'VISHWAKARMA INSTITUTE OF TECHNOLOGY, PUNE – 411037.\n'
        '(An Autonomous Institute Affiliated to Savitribai Phule Pune University)'
    )
    doc.add_paragraph('Examination: ESE')
    doc.add_paragraph('Year:\t\tBranch: AI&DS')
    doc.add_paragraph('Subject:\t\t\tSubject Code:')
    doc.add_paragraph('Max. Marks:\t\t\tTotal Pages of Question Paper:')
    doc.add_paragraph('Day & Date:\t\t\tTime:')
    doc.add_paragraph('¬¬¬¬¬¬¬¬¬¬¬¬¬¬¬¬¬¬¬¬¬')

    # Instructions to candidates
    doc.add_paragraph('Instructions to Candidate')
    doc.add_paragraph('1. All questions are compulsory.')
    doc.add_paragraph('2. Neat diagrams must be drawn wherever necessary.')
    doc.add_paragraph('3. Figures to the right indicate full marks.')

    # Create the table for questions
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Fill header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Q.N.'
    hdr_cells[1].text = 'BT* No'
    hdr_cells[2].text = 'Max marks'

    # Add questions to the table
    question_count = 1
    sub_question_count = 0
    bt_levels = [1, 2, 3, 4, 5, 6]  # Bloom's Taxonomy levels

    for question in questions:
        if sub_question_count == 0:
            # Add a main question header row
            main_row = table.add_row().cells
            main_row[0].text = f"Q.{question_count}."
            main_row[1].text = ''
            main_row[2].text = '10'

        # Add a sub-question row
        sub_question_count += 1
        row = table.add_row().cells
        row[0].text = f"({chr(96 + sub_question_count)})"  # a, b, c...
        row[1].text = question['question']  # Insert question text
        row[2].text = str(bt_levels[(question_count - 1) % len(bt_levels)])  # Assign BT level cyclically

        if sub_question_count == 3:
            sub_question_count = 0
            question_count += 1

    # CO statements section
    doc.add_paragraph('\nCO Statements:')
    co_statements = [
        "CO1: Demonstrate knowledge learning algorithms and concept learning.",
        "CO2: Evaluate decision tree learning algorithm.",
        "CO3: Formulate a given problem within the Bayesian learning framework and SVM.",
        "CO4: Apply different clustering algorithms used in machine learning.",
        "CO5: Explore Association rule mining and dimensionality reduction.",
        "CO6: Analyze research-based problems using machine learning techniques like Reinforcement learning and ANN."
    ]
    for statement in co_statements:
        doc.add_paragraph(statement)

    # Blooms Taxonomy Levels
    doc.add_paragraph('\n*Blooms Taxonomy (BT) Level No:')
    bt_levels = [
        "1. Remembering; 2. Understanding; 3. Applying; 4. Analyzing; 5. Evaluating; 6. Creating"
    ]
    for level in bt_levels:
        doc.add_paragraph(level)

    # Save the document
    doc.save(output_path)
    print("Document structure created successfully.")

# Main Function
def main(questions):
    # Specify the output file path
    output_structure_path = 'question_paper_generated.docx'
    create_question_paper_structure(output_structure_path, questions)

if __name__ == "__main__":
    # Example questions (replace with actual questions from the API)
    example_questions = [
        {"qno": 1, "question": "Explain the concept of machine learning and provide an example.", "options": [], "answer": "Machine learning is a subset of AI that enables systems to learn from data."},
        {"qno": 2, "question": "Discuss the challenges in concept learning.", "options": [], "answer": "Challenges include overfitting, underfitting, and data quality."},
        {"qno": 3, "question": "Describe a situation where concept learning is beneficial.", "options": [], "answer": "Concept learning is beneficial in spam detection."}
    ]
    main(example_questions)